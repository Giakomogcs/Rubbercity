#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Conversor Markdown -> DOCX com identidade visual SENAI.
Gera documentacao/Rubbercity.docx a partir de documentacao/DOCUMENTACAO.md.

Recursos: capa estilizada, header/footer com numero de pagina, sumario (TOC),
titulos coloridos, tabelas zebradas, callouts (NOTE/TIP/WARNING/DANGER),
passos numerados com badge, blocos de codigo escuros, markdown inline,
e diagramas Mermaid renderizados em imagem (via mmdc) com fallback para bloco.
"""
import os
import re
import shutil
import subprocess
import tempfile
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- Paleta SENAI ----------
ACCENT   = "E30613"
ACCENT2  = "8B0410"
INK      = "0B0B0B"
INK2     = "1F2937"
GRAY     = "6B7280"
RULE     = "D1D5DB"
ZEBRA    = "F9F7F7"
CELLHDR  = "FBE9EB"
CODEBG   = "0B0B0B"
CODEFG   = "E5E7EB"
DANGER   = "B91C1C"
WARN     = "B45309"
TIP      = "047857"
INFO     = "1D4ED8"

CALLOUTS = {
    "NOTE":    ("Nota",      INFO,   "F0F5FF", "i"),
    "TIP":     ("Dica",      TIP,    "ECFDF5", "+"),
    "WARNING": ("Atencao",   WARN,   "FFFBEB", "!"),
    "DANGER":  ("Cuidado",   DANGER, "FEF2F2", "x"),
}

BASE = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(BASE, "DOCUMENTACAO.md")
OUT_PATH = os.path.join(BASE, "Rubbercity.docx")
IMG_DIR = os.path.join(BASE, "_diagrams")

PROJECT = "Rubbercity — Agente OFM + RAG"
DOC_TITLE = "Documentacao do Projeto"


def rgb(hexstr):
    return RGBColor.from_string(hexstr)


def shade(el, fill):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), fill)
    el.append(sh)


def set_cell_bg(cell, fill):
    shade(cell._tc.get_or_add_tcPr(), fill)


def set_cell_borders(cell, color=RULE, sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)


def set_cell_margins(cell, top=60, bottom=60, left=110, right=110):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def left_bar(cell, color):
    """Barra lateral grossa colorida (para callouts)."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    e = OxmlElement("w:left")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), "24")
    e.set(qn("w:space"), "0")
    e.set(qn("w:color"), color)
    borders.append(e)
    tcPr.append(borders)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")


def add_inline(par, text, base_color=INK2, base_size=10.5, base_bold=False):
    """Renderiza **negrito**, *italico*, `codigo` inline."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = par.add_run(part[2:-2]); r.bold = True
            r.font.color.rgb = rgb(base_color); r.font.size = Pt(base_size)
        elif part.startswith("`") and part.endswith("`"):
            r = par.add_run(part[1:-1]); r.font.name = "Consolas"
            r.font.size = Pt(base_size - 0.5); r.font.color.rgb = rgb(ACCENT2)
            shade(r._r.get_or_add_rPr(), "F3F4F6")
        elif part.startswith("*") and part.endswith("*"):
            r = par.add_run(part[1:-1]); r.italic = True
            r.font.color.rgb = rgb(base_color); r.font.size = Pt(base_size)
        else:
            r = par.add_run(part)
            r.font.color.rgb = rgb(base_color); r.font.size = Pt(base_size)
            r.bold = base_bold


def add_field(par, instr):
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate")
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
    r = par.add_run()
    for el in (fb, it, fs):
        r._r.append(el)
    r2 = par.add_run("•"); r2._r.append(fe)


# ---------- Mermaid ----------
def find_mmdc():
    for cand in ("mmdc", "mmdc.cmd"):
        if shutil.which(cand):
            return shutil.which(cand)
    return None


def render_mermaid(code, idx):
    mmdc = find_mmdc()
    if not mmdc:
        return None
    os.makedirs(IMG_DIR, exist_ok=True)
    out = os.path.join(IMG_DIR, f"diagram_{idx}.png")
    with tempfile.NamedTemporaryFile("w", suffix=".mmd", delete=False, encoding="utf-8") as f:
        f.write(code); tmp = f.name
    try:
        subprocess.run([mmdc, "-i", tmp, "-o", out, "-b", "white", "-s", "2"],
                       check=True, capture_output=True, timeout=120)
        return out if os.path.exists(out) else None
    except Exception:
        return None
    finally:
        try:
            os.unlink(tmp)
        except OSError:
            pass


# ---------- Estilos base ----------
def setup_base_styles(doc):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.font.color.rgb = rgb(INK2)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.15


def heading(doc, text, level):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if level == 1:
        pf.space_before = Pt(18); pf.space_after = Pt(8)
        r = p.add_run(text.upper()); r.bold = True
        r.font.size = Pt(18); r.font.color.rgb = rgb(ACCENT)
        r.font.name = "Calibri"
        bottom_border(p, ACCENT, 12)
    elif level == 2:
        pf.space_before = Pt(14); pf.space_after = Pt(6)
        r = p.add_run(text); r.bold = True
        r.font.size = Pt(14.5); r.font.color.rgb = rgb(INK)
        bottom_border(p, RULE, 6)
    elif level == 3:
        pf.space_before = Pt(10); pf.space_after = Pt(4)
        r = p.add_run(text); r.bold = True
        r.font.size = Pt(12); r.font.color.rgb = rgb(INK2)
    else:
        pf.space_before = Pt(8); pf.space_after = Pt(2)
        r = p.add_run(text); r.bold = True
        r.font.size = Pt(11); r.font.color.rgb = rgb(GRAY)
    return p


def bottom_border(par, color, sz):
    pPr = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), str(sz))
    b.set(qn("w:space"), "2"); b.set(qn("w:color"), color)
    pbdr.append(b); pPr.append(pbdr)


# ---------- Capa ----------
def build_cover(doc):
    # faixa superior vermelha
    band = doc.add_paragraph()
    band.paragraph_format.space_before = Pt(36); band.paragraph_format.space_after = Pt(0)
    rb = band.add_run("SENAI"); rb.bold = True; rb.font.size = Pt(13)
    rb.font.color.rgb = rgb("FFFFFF")
    pPr = band._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), ACCENT); pPr.append(shd)
    band.paragraph_format.left_indent = Pt(-2)

    for _ in range(5):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = t.add_run(DOC_TITLE.upper()); r.bold = True
    r.font.size = Pt(34); r.font.color.rgb = rgb(ACCENT)
    bottom_border(t, ACCENT, 18)

    s = doc.add_paragraph()
    r = s.add_run(PROJECT); r.font.size = Pt(18); r.font.color.rgb = rgb(INK)
    r.bold = True

    sub = doc.add_paragraph()
    r = sub.add_run("Documentacao de Negocio e Tecnica — Engenharia Reversa do Codigo-Fonte")
    r.font.size = Pt(12); r.italic = True; r.font.color.rgb = rgb(GRAY)

    for _ in range(10):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    r = meta.add_run(f"Data: {date.today().strftime('%d/%m/%Y')}\n")
    r.font.size = Pt(11); r.font.color.rgb = rgb(INK2)
    r = meta.add_run("Autor: Doc Master\n"); r.font.size = Pt(11); r.font.color.rgb = rgb(INK2)
    r = meta.add_run("Confidencial — uso interno"); r.font.size = Pt(10); r.font.color.rgb = rgb(ACCENT2)
    bottom_border(meta, ACCENT2, 8)

    doc.add_page_break()


# ---------- TOC ----------
def build_toc(doc):
    h = doc.add_paragraph()
    r = h.add_run("SUMARIO"); r.bold = True; r.font.size = Pt(16)
    r.font.color.rgb = rgb(ACCENT)
    bottom_border(h, ACCENT, 10)
    p = doc.add_paragraph()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = r'TOC \o "1-3" \h \z \u'
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate")
    ft = OxmlElement("w:t"); ft.text = "Atualize o campo (F9) para gerar o sumario."
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
    run = p.add_run()
    for el in (fb, it, fs, ft, fe):
        run._r.append(el)
    doc.add_page_break()


# ---------- Header / Footer ----------
def build_header_footer(doc):
    sec = doc.sections[-1]
    hdr = sec.header
    hp = hdr.paragraphs[0]
    hp.text = ""
    tab = hp.paragraph_format
    r = hp.add_run(DOC_TITLE + "  ")
    r.font.size = Pt(8); r.font.color.rgb = rgb(GRAY)
    hp.add_run("\t")
    r = hp.add_run(PROJECT); r.font.size = Pt(8); r.font.color.rgb = rgb(ACCENT)
    bottom_border(hp, RULE, 4)
    _set_tabs(hp)

    ftr = sec.footer
    fp = ftr.paragraphs[0]
    fp.text = ""
    r = fp.add_run("SENAI  ")
    r.font.size = Pt(8); r.font.color.rgb = rgb(ACCENT); r.bold = True
    fp.add_run("\t")
    fp.add_run("Pagina ")
    add_page_field(fp, "PAGE")
    fp.add_run(" de ")
    add_page_field(fp, "NUMPAGES")
    for run in fp.runs:
        run.font.size = Pt(8); 
        if run.font.color.rgb is None:
            run.font.color.rgb = rgb(GRAY)
    _set_tabs(fp)


def add_page_field(par, instr):
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
    r = par.add_run()
    r._r.append(fb); r._r.append(it); r._r.append(fe)
    r.font.size = Pt(8); r.font.color.rgb = rgb(GRAY)


def _set_tabs(par):
    pPr = par._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    t = OxmlElement("w:tab")
    t.set(qn("w:val"), "right"); t.set(qn("w:pos"), "9360")
    tabs.append(t); pPr.append(tabs)


# ---------- Conteudo ----------
def add_paragraph_text(doc, text):
    p = doc.add_paragraph()
    add_inline(p, text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    add_inline(p, text)
    return p


def add_numbered_step(doc, num, text):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    badge = table.cell(0, 0)
    body = table.cell(0, 1)
    badge.width = Inches(0.35); body.width = Inches(6.3)
    set_cell_bg(badge, ACCENT)
    bp = badge.paragraphs[0]; bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = bp.add_run(str(num)); r.bold = True; r.font.color.rgb = rgb("FFFFFF"); r.font.size = Pt(10)
    set_cell_margins(badge, 30, 30, 20, 20)
    bodyp = body.paragraphs[0]
    add_inline(bodyp, text)
    set_cell_margins(body, 30, 30, 110, 60)
    _no_table_borders(table)


def _no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, CODEBG)
    set_cell_margins(cell, 90, 90, 140, 140)
    cell.paragraphs[0].text = ""
    for i, ln in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0
        r = p.add_run(ln if ln else " ")
        r.font.name = "Consolas"; r.font.size = Pt(9); r.font.color.rgb = rgb(CODEFG)
    _no_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_callout(doc, kind, lines):
    label, color, fill, icon = CALLOUTS[kind]
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, fill)
    left_bar(cell, color)
    set_cell_margins(cell, 80, 80, 160, 120)
    p0 = cell.paragraphs[0]
    r = p0.add_run(f"{icon}  {label.upper()}")
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = rgb(color)
    body = cell.add_paragraph()
    add_inline(body, " ".join(lines), base_color=INK2, base_size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table_block(doc, rows):
    header = [c.strip() for c in rows[0]]
    data = rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, htext in enumerate(header):
        set_cell_bg(hdr[i], ACCENT)
        set_cell_borders(hdr[i], RULE)
        set_cell_margins(hdr[i])
        p = hdr[i].paragraphs[0]
        add_inline(p, htext, base_color="FFFFFF", base_size=10, base_bold=True)
        for run in p.runs:
            run.bold = True; run.font.color.rgb = rgb("FFFFFF")
    for ri, row in enumerate(data):
        cells = table.add_row().cells
        fill = ZEBRA if ri % 2 == 0 else "FFFFFF"
        for ci in range(len(header)):
            val = row[ci].strip() if ci < len(row) else ""
            set_cell_bg(cells[ci], fill)
            set_cell_borders(cells[ci], RULE)
            set_cell_margins(cells[ci])
            add_inline(cells[ci].paragraphs[0], val, base_size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_table(block_lines):
    rows = []
    for ln in block_lines:
        if re.match(r"^\s*\|?\s*:?-{2,}", ln):  # separador
            continue
        cells = [c for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


# ---------- Loop principal ----------
def convert():
    with open(MD_PATH, encoding="utf-8") as f:
        md = f.read()
    lines = md.split("\n")

    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.9); s.bottom_margin = Inches(0.9)
        s.left_margin = Inches(0.9); s.right_margin = Inches(0.9)
    setup_base_styles(doc)
    build_cover(doc)
    build_toc(doc)
    build_header_footer(doc)

    i = 0
    diag_idx = 0
    n = len(lines)
    skip_h1_title = True  # primeiro H1 do MD vira capa
    while i < n:
        ln = lines[i]
        stripped = ln.strip()

        # Codigo / mermaid
        if stripped.startswith("```"):
            lang = stripped[3:].strip().lower()
            block = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                block.append(lines[i]); i += 1
            i += 1
            if lang == "mermaid":
                diag_idx += 1
                img = render_mermaid("\n".join(block), diag_idx)
                if img:
                    pic = doc.add_paragraph(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pic.add_run().add_picture(img, width=Inches(6.2))
                    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    rc = cap.add_run(f"Figura {diag_idx} — diagrama do sistema")
                    rc.italic = True; rc.font.size = Pt(9); rc.font.color.rgb = rgb(GRAY)
                else:
                    add_code_block(doc, block)
                    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    rc = cap.add_run(f"(diagrama Mermaid — Figura {diag_idx})")
                    rc.italic = True; rc.font.size = Pt(9); rc.font.color.rgb = rgb(GRAY)
            else:
                add_code_block(doc, block)
            continue

        # Callout
        m = re.match(r"^>\s*\[!(NOTE|TIP|WARNING|DANGER)\]", stripped)
        if m:
            kind = m.group(1)
            body = []
            i += 1
            while i < n and lines[i].strip().startswith(">"):
                body.append(lines[i].strip().lstrip(">").strip())
                i += 1
            add_callout(doc, kind, [b for b in body if b])
            continue

        # Tabela
        if stripped.startswith("|") and i + 1 < n and re.match(r"^\s*\|?\s*:?-{2,}", lines[i + 1]):
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i]); i += 1
            add_table_block(doc, parse_table(block))
            continue

        # Headings
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            if level == 1 and skip_h1_title:
                skip_h1_title = False
                i += 1
                continue
            heading(doc, text, level)
            i += 1
            continue

        # HR
        if stripped == "---":
            i += 1
            continue

        # Numbered list
        mnum = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if mnum:
            add_numbered_step(doc, mnum.group(1), mnum.group(2))
            i += 1
            continue

        # Bullet
        if stripped.startswith("- "):
            add_bullet(doc, stripped[2:])
            i += 1
            continue

        # Paragrafo / vazio
        if stripped:
            add_paragraph_text(doc, stripped)
        i += 1

    doc.save(OUT_PATH)
    print(f"OK: {OUT_PATH}")
    print(f"Mermaid CLI: {'encontrado' if find_mmdc() else 'NAO encontrado (diagramas como bloco de codigo)'}")


if __name__ == "__main__":
    convert()
